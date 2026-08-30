#!/usr/bin/env python3
"""Converte os documentos jurídicos em Markdown para .docx formatado.

Suporta o subconjunto de Markdown usado em docs/juridico/: cabeçalhos (# a ###),
parágrafos com **negrito**, tabelas com <br> nas células, listas, citações e
réguas horizontais. Uso: md2docx.py entrada.md saida.docx "Título do documento"
"""
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

GOLD = RGBColor(0xC9, 0xA8, 0x4C)
INK = RGBColor(0x0A, 0x0A, 0x0A)


def setup(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def footer_page_numbers(doc):
    """Rodapé com o campo PAGE (numeração automática do Word)."""
    par = doc.sections[0].footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    run.font.size = Pt(8)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)


def shade(cell, color):
    tc = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), color)
    tc.append(el)


def add_runs(par, text, bold=False):
    """Renderiza **negrito** e `código` dentro de um parágrafo."""
    text = text.replace("<br>", "\n")
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        strong = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if strong else part
        for i, line in enumerate(content.split("\n")):
            run = par.add_run(line)
            run.bold = bold or strong
            if i < len(content.split("\n")) - 1:
                run.add_break()


def add_table(doc, rows):
    header, body = rows[0], rows[2:]
    table = doc.add_table(rows=0, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, cells in enumerate([header] + body):
        row = table.add_row().cells
        for cell, text in zip(row, cells + [""] * (len(header) - len(cells))):
            cell.paragraphs[0].text = ""
            par = cell.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.space_after = Pt(2)
            add_runs(par, text, bold=(idx == 0))
            for run in par.runs:
                run.font.size = Pt(9)
            if idx == 0:
                shade(cell, "1C1C1C")
                for run in par.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_path, docx_path):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    setup(doc)
    footer_page_numbers(doc)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "<br>":
            i += 1
            continue

        if stripped == "---":
            par = doc.add_paragraph()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "C9A84C")
            pbdr.append(bottom)
            par._p.get_or_add_pPr().append(pbdr)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].strip().replace("|", "").replace(" ", "")
        ) <= {"-", ":"} and lines[i + 1].strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(split_row(lines[i]))
                i += 1
            add_table(doc, block)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("# ").strip()
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            par.paragraph_format.space_after = Pt(8)
            par.paragraph_format.keep_with_next = True
            add_runs(par, text, bold=True)
            for run in par.runs:
                run.font.size = Pt({1: 15, 2: 12.5, 3: 11}.get(level, 10.5))
                run.font.color.rgb = INK if level == 1 else GOLD if level == 2 else INK
            i += 1
            continue

        if stripped.startswith("> "):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.8)
            add_runs(par, stripped[2:])
            for run in par.runs:
                run.italic = True
                run.font.size = Pt(9.5)
            i += 1
            continue

        if re.match(r"^[-*] ", stripped) or re.match(r"^\d+\. ", stripped):
            bullet = re.match(r"^[-*] ", stripped)
            text = re.sub(r"^([-*]|\d+\.) ", "", stripped)
            par = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            par.paragraph_format.space_after = Pt(3)
            add_runs(par, text)
            i += 1
            continue

        par = doc.add_paragraph()
        add_runs(par, stripped)
        i += 1

    doc.save(docx_path)
    print(f"gerado: {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
